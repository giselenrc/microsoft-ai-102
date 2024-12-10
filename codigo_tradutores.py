from typing import ParamSpec
import requests
from docx import Document
import os

subscription_key = "c39907db83554d219aed77c52f568c67"
endpoint = "https://api.cognitive.microsofttranslator.com"
location = "eastus2"
target_language = 'pt-br'


def translator_text(text, target_language):
    path = '/translate'
    constructed_url = endpoint + path

    headers = {
        'Ocp-Apim-Subscription-Key': subscription_key,
        'Ocp-Apim-Subscription-Region': location,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(os.urandom(16))
    }

    body = [
        {
            "text": text  # Changed the structure of the body
        }
    ]

    params = {
        'api-version': '3.0',
        'from': 'en',
        'to': target_language
    }

    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()

    # Check if the response contains an error
    if 'error' in response:
        print(f"Error: {response['error']}")
        return None  # or raise an exception

    # Assuming the response structure is correct if no error
    # Adjust this part if the structure differs based on the API response
    try:
        return response[0]["translations"][0]["text"]
    except (IndexError, KeyError):
        print(f"Unexpected response format: {response}")
        return None  # or raise an exception


def translate_document(path):
    original_document = Document(path) 
    translated_doc = Document()
    
    for paragraph in original_document.paragraphs:
        translated_text = translator_text(paragraph.text, target_language) 
        translated_doc.add_paragraph(translated_text)

    path_translated = path.replace(".docx", f"_{target_language}.docx")
    translated_doc.save(path_translated)

    return translated_text
    